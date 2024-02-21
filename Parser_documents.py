"""# Importar las bibliotecas necesarias
from docx import Document
import pdfplumber

class Parser_Documents:
    def __init__(self, file):
        self.file = file

    def parse(self):
        if self.file.filename.endswith('.txt'):
            return self.parse_text()
        elif self.file.filename.endswith('.docx'):
            return self.parse_docx()
        elif self.file.filename.endswith('.pdf'):
            return self.parse_pdf()
        else:
            raise ValueError("Archivo no soportado")

    def parse_text(self):
        with open(source, 'r', encoding='utf-8') as file:
            content = file.read()
            return content
        #return self.file.read().decode('utf-8')

    def parse_docx(self):
        doc = Document(self.file)
        return "\n".join([para.text for para in doc.paragraphs])

    def parse_pdf(self):
        with pdfplumber.open(self.file) as pdf:
            text = ''
            for page in pdf.pages:
                text += page.extract_text() + '\n'
            return text """

import pdfplumber
from docx import Document

class Parser_Documents:
    def __init__(self, file):
        self.file = file

    def parse(self):
        """

        :rtype: object
        """
        # Si el objeto 'file' es un string, asumimos que es una ruta de archivo local
        if isinstance(self.file, str):
            file_extension = self.file.split('.')[-1].lower()
            with open(self.file, 'r', encoding='utf-8') as file:
                content = file.read()
                return content

        # Para objetos de archivo (stream), verificar la extensión y procesar
        if hasattr(self.file, 'filename'):
            if self.file.filename.endswith('.txt'):
                self.file.seek(0)  # Reiniciar el puntero del archivo
                content = self.file.read().decode('utf-8')

                return content
            elif self.file.filename.endswith('.docx'):
                return self.parse_docx()
            elif self.file.filename.endswith('.pdf'):
                return self.parse_pdf()
            else:
                raise ValueError("Unsupported file format")

    def parse_docx(self):
        doc = Document(self.file)
        return "\n".join([para.text for para in doc.paragraphs])

    def parse_pdf(self):
        with pdfplumber.open(self.file) as pdf:
            text = ''
            for page in pdf.pages:
                text += page.extract_text() + '\n'
            return text